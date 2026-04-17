import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- SETTINGS ---
OUT_FILE = "LME_Analytics_Logic_Guide.docx"
UM6P_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
UM6P_ORANGE = RGBColor(0xF4, 0x79, 0x20)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = UM6P_BLUE
    return h

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def generate_doc():
    doc = Document()

    # --- TITLE PAGE ---
    doc.add_paragraph("\n" * 5)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("LME Analytics Logic & Computation Guide")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = UM6P_BLUE

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = st.add_run("State Triggers, Risk Formulas, and Testing Troubleshooting")
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = UM6P_ORANGE

    doc.add_paragraph("\n" * 10)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("April 2026\nUniversité Mohammed VI Polytechnique")
    run.font.size = Pt(12)
    doc.add_page_break()

    # --- SECTION 1: STATE TRIGGERS ---
    add_styled_heading(doc, "1. Real-Time Learner States", 1)
    doc.add_paragraph("States update every second. They classify the student's immediate mental state based on their behavior.")
    
    doc.add_paragraph("• Engaged: Default state. Reset on every interaction or correct answer.")
    doc.add_paragraph("• Struggling: Triggered if (Idle_Time >= 30s) OR (Attempts >= 3). This indicates high cognitive load.")
    doc.add_paragraph("• Unengaged: Triggered if (Idle_Time > 60s). This indicates a behavioral dropout or loss of focus.")
    
    doc.add_page_break()

    # --- SECTION 2: CORE MATHEMATICAL FORMULAS ---
    add_styled_heading(doc, "2. Core Mathematical Formulas", 1)
    
    add_styled_heading(doc, "2.1 Struggle Index (SI)", 2)
    add_formula(doc, "SI = [0.4 * (Idle / 120)] + [0.4 * (Attempts / 5)] - [0.2 * (Score / 1)]")
    doc.add_paragraph("The SI measures task difficulty. It increases with idle time and mistakes, but is lowered by a high mastery score.")

    add_styled_heading(doc, "2.2 Mastery Trajectory (EMA)", 2)
    add_formula(doc, "Mastery_n = (0.3 * Current_Score) + (0.7 * Previous_Mastery)")
    doc.add_paragraph("Instead of a simple average, we use an Exponential Moving Average (EMA) with Alpha=0.3 to ensure the dashboard prioritizes the most recent learning consolidation.")

    add_styled_heading(doc, "2.3 Global Risk Score", 2)
    add_formula(doc, "Risk = (0.3 * Avg_SI) + (0.3 * Hint_Dep) + (0.2 * Unengaged_Ratio) + (0.2 * (1 - Part_Rate))")
    doc.add_paragraph("The Risk Score is the final diagnostic. The 'Participation Penalty' ensures that students who skip questions are flagged.")
    
    doc.add_page_break()

    # --- SECTION 3: DIAGNOSTIC PRIORITY & LABELS ---
    add_styled_heading(doc, "3. Risk Labels & Diagnostic Priority", 1)
    
    doc.add_paragraph("Risk Labels (Current Thresholds):")
    doc.add_paragraph("• Low Risk: Score < 0.30")
    doc.add_paragraph("• Medium Risk: Score 0.30 - 0.50")
    doc.add_paragraph("• High Risk: Score > 0.50")

    add_styled_heading(doc, "3.1 Weakness Priority Logic", 2)
    doc.add_paragraph("The system checks for weaknesses in this strict order to capture the most critical issues first:")
    doc.add_paragraph("1. Emotional: Triggered by Frustration (rapid mistakes) or Boredom (unengaged ratio).")
    doc.add_paragraph("2. Behavioral: Triggered by Impulsivity (high attempts) or Time Abandonment (avg idle > 2min).")
    doc.add_paragraph("3. Cognitive: Default for high hints or a rising struggle index.")
    
    doc.add_page_break()

    # --- SECTION 4: TESTING EDGE CASES & TROUBLESHOOTING ---
    add_styled_heading(doc, "4. Troubleshooting Your Demo Results", 1)
    
    add_styled_heading(doc, "4.1 Why High Risk might not appear", 2)
    doc.add_paragraph("The most common reason is 'Session Dilution.' If a student has 2 'perfect' sessions in the database, the system will average them with the new session. A single bad session might only move the student from 'Low' to 'Medium'.")
    doc.add_paragraph("Solution: Run 'python clear_adam.py' to reset the student history before testing a specific high-risk scenario.")

    add_styled_heading(doc, "4.2 Why 'Unengaged' didn't trigger", 2)
    doc.add_paragraph("If you skip questions very fast (e.g., every 5 seconds), you never reach the 60-second idle threshold required for 'Unengaged.'")
    doc.add_paragraph("However, the system still catches this as a risk because 'Participation Rate' will drop to 10%, triggering the 0.2 penalty factor in the Risk Score.")

    add_styled_heading(doc, "4.3 Triggering Emotional (Frustration)", 2)
    doc.add_paragraph("Mistakes must be FAST (<10s apart) to trigger the 'Frustration Index.' If you wait too long between clicking wrong answers, the system classifies it as 'Behavioral' trial-and-error.")

    # --- SAVE ---
    doc.save(OUT_FILE)
    print(f"Generated {OUT_FILE}")

if __name__ == "__main__":
    generate_doc()
