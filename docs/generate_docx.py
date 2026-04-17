import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CONFIGURATION ---
OUT_FILE = "LME_Analytics_Current_State.docx"
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(DOCS_DIR, OUT_FILE)

# --- THEME COLORS ---
UM6P_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
UM6P_ORANGE = RGBColor(0xF4, 0x79, 0x20)
HEADER_BG = RGBColor(0x1E, 0x3A, 0x5F)
ROW_ALT_BG = RGBColor(0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LME Analytics Layer — Confidential · UM6P 2026")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = UM6P_BLUE
            run.font.size = Pt(16)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)

    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], HEADER_BG)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row().cells
        for c_idx, val in enumerate(row_data):
            row[c_idx].text = str(val)
            if r_idx % 2 == 1:
                set_cell_bg(row[c_idx], ROW_ALT_BG)
            p = row[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.size = Pt(9)
    return table

# --- MAIN GENERATION ---
doc = Document()

# 1. TITLE PAGE
doc.add_paragraph("\n" * 5)
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("LME Analytics Layer — Current State Documentation")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = UM6P_BLUE

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = st.add_run("Learner Profiling System · UM6P · 2026")
run.font.size = Pt(16)
run.font.italic = True
run.font.color.rgb = UM6P_ORANGE

doc.add_paragraph("\n" * 10)
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("April 2026\nUniversité Mohammed VI Polytechnique")
run.font.size = Pt(12)

doc.add_page_break()

# 2. TABLE OF CONTENTS
add_heading(doc, "Table of Contents", 1)
doc.add_paragraph("Section 1 — System Overview")
doc.add_paragraph("Section 2 — Data Foundation")
doc.add_paragraph("Section 3 — Metric Definitions (by Fredricks Dimension)")
doc.add_paragraph("Section 4 — Risk Model")
doc.add_paragraph("Section 5 — API Endpoints")
doc.add_paragraph("Section 6 — Test Scenarios")
doc.add_paragraph("Section 7 — Known Limitations and Future Work")
doc.add_page_break()

# SECTION 1
add_heading(doc, "Section 1 — System Overview", 1)
doc.add_paragraph(
    "The LME Analytics Layer is a secondary diagnostic system that sits passively on top of the existing Learner Modeling Engine. "
    "Its primary purpose is to transform raw interaction telemetry into actionable pedagogical insights for instructors, without "
    "interfering with the core quiz engine or database schema."
)
doc.add_paragraph(
    "The layer consists of three primary components:\n"
    "1. seed_analytics.py: Responsible for generating high-fidelity synthetic data for simulation and testing.\n"
    "2. risk_model.py: Implements the machine learning and fallback logic for student risk assessment.\n"
    "3. analytics.py: The entry point providing REST endpoints and real-time metric computation."
)
doc.add_paragraph(
    "Theoretical Foundation: The system is built upon the engagement framework by Fredricks et al. (2004), "
    "which classifies learner behavior across three distinct dimensions:\n"
    "• Behavioral: Observable actions like participation, idle time, and attempts.\n"
    "• Cognitive: Depth of processing evidenced by mastery trajectory and hint dependency.\n"
    "• Emotional: Affective states and resilience, operationalized through frustration and recovery patterns."
)
doc.add_page_break()

# SECTION 2
add_heading(doc, "Section 2 — Data Foundation", 1)
add_heading(doc, "2.1 Telemetry Logs Table Structure", 2)
headers = ["Column", "Type", "Description"]
rows = [
    ["id", "INTEGER", "Primary Key"],
    ["student_id", "STRING", "Unique identifier for the student"],
    ["module", "STRING", "Module name (maths/logic)"],
    ["current_question_text", "STRING", "The text of the active question"],
    ["idle_time", "FLOAT", "Seconds of inactivity on current question"],
    ["attempt_count", "INTEGER", "Total attempts on current question"],
    ["is_correct", "BOOLEAN", "Result of the current attempt"],
    ["current_score", "FLOAT", "Running mastery score [0-1]"],
    ["total_mistakes", "INTEGER", "Cumulative errors in session"],
    ["total_idle_time", "FLOAT", "Cumulative idle time in session"],
    ["answered_count", "INTEGER", "Number of questions answered"],
    ["total_hints_requested", "INTEGER", "Cumulative hints used"],
    ["learner_state", "STRING", "Classification: Engaged, Struggling, Unengaged"],
    ["struggle_index", "FLOAT", "Derived complexity metric"],
    ["timestamp", "DATETIME", "ISO 8601 recorded time"],
]
add_table(doc, headers, rows)

add_heading(doc, "2.2 Session Boundary Logic (3-Rule System)", 2)
add_code_block(doc, 
"was_completed = df['type'].shift() == 'session_complete'\ntime_gap = df['datetime'].diff() > pd.Timedelta(minutes=5)\nmodule_change = df['module'] != df['module'].shift()\n\ndf['session_index'] = (was_completed | time_gap | module_change).cumsum() + 1")

add_heading(doc, "2.3 Struggle Index Formula", 2)
add_code_block(doc, "SI = (0.4 * (idle_time / 120.0)) + (0.4 * (attempt_count / 5.0)) - (0.2 * (current_score / 1.0))")

add_heading(doc, "2.4 Learner State Classification Rules", 2)
add_code_block(doc, 
"if is_correct: 'Engaged'\nelif attempt_count == 0 and idle_time < 20: 'Engaged'\nelif idle_time > 60: 'Unengaged'\nelif idle_time >= 30 or attempt_count >= 3: 'Struggling'\nelse: 'Engaged'")
doc.add_page_break()

# SECTION 3
add_heading(doc, "Section 3 — Metric Definitions (by Fredricks Dimension)", 1)
add_heading(doc, "3.1 Behavioral Dimension", 2)
rows = [
    ["Avg Idle Time", "Behavioral", "mean(total_idle_time per session)", "Measures physical presence and focus intensity."],
    ["Avg Attempt Count", "Behavioral", "mean(attempt_count per row)", "Measures persistence vs. impulsive guessing."],
    ["Participation Rate", "Behavioral", "answered_count / questions_in_module", "Measures overall session completion."],
    ["Total Mistakes", "Behavioral", "sum(is_correct == False)", "Measures error frequency."],
]
add_table(doc, ["Metric", "Dimension", "Logic", "Description"], rows)

add_heading(doc, "3.2 Cognitive Dimension", 2)
rows = [
    ["Mastery (EMA)", "Cognitive", "EMA(current_score, alpha=0.3)", "Smoothed growth trajectory prioritizing recent sessions."],
    ["Struggle Index Trend", "Cognitive", "mean(SI per session)", "Measures difficulty progression over time."],
    ["Hint Dependency", "Cognitive", "total_hints / max(1, answered_count)", "Measures reliance on external scaffolding."],
    ["Avg Hint Level", "Cognitive", "mean(hint_dependency list)", "Measures sophistication of help-seeking."],
]
add_table(doc, ["Metric", "Dimension", "Logic", "Description"], rows)

add_heading(doc, "3.3 Emotional Dimension", 2)
rows = [
    ["State Ratios", "Emotional", "count(state) / total_rows", "Balance between engagement, struggle, and boredom."],
    ["Recovery Events", "Emotional", "Count cases where Struggling -> Engaged via hintless correct answer", "Measures resilience and self-correction capability."],
    ["Frustration Index", "Emotional", "count(attempts>=3 & incorrect & idle<10) / rows", "Measures impulsive repeated errors due to frustration."],
]
add_table(doc, ["Metric", "Dimension", "Logic", "Description"], rows)
doc.add_page_break()

# SECTION 4
add_heading(doc, "Section 4 — Risk Model", 1)
doc.add_paragraph("Model: Logistic Regression (max_iter=500). Fallback: Deterministic weighted sum (0.4*SI + 0.4*Hint + 0.2*Unengaged).")
rows = [
    ["avg_SI", "student_df['struggle_index'].mean()", "Cognitive"],
    ["hint_dependency_score", "global_hints / global_answered", "Cognitive"],
    ["unengaged_ratio", "unengaged_rows / total_rows", "Emotional"],
    ["avg_attempt_count", "student_df['attempt_count'].mean()", "Behavioral"],
    ["session_count", "len(sessions)", "Behavioral"],
]
add_table(doc, ["Feature", "Formula", "Dimension"], rows)

add_heading(doc, "Dominant Weakness Priority", 2)
add_code_block(doc, "if is_behav: 'behavioral'\nelif is_cog: 'cognitive'\nelif is_emot: 'emotional'\nelse: 'cognitive'")
doc.add_page_break()

# SECTION 5
add_heading(doc, "Section 5 — API Endpoints", 1)
doc.add_paragraph("Endpoints require valid Instructor JWT in headers.")
rows = [
    ["GET /analytics/student/{id}", "Instructor", "Dimensional profile for student", "Full JSON with Fredricks buckets"],
    ["GET /analytics/cohort", "Instructor", "Aggregated class metrics", "Class distribution & ranking"],
    ["GET /analytics/risk", "Instructor", "ML predictions & interventions", "Risk Score & Weakness mapping"],
    ["GET /analytics/seed", "Open", "Wipe and reseed DB", "Subprocess execution status"],
]
add_table(doc, ["Path", "Auth", "Function", "Result"], rows)
doc.add_page_break()

# SECTION 6
add_heading(doc, "Section 6 — Test Scenarios", 1)
headers = ["Scenario", "Profile", "Conditions", "SI Range", "Risk", "Weakness", "Dashboard / Intervention"]
scenarios = [
    ["1. Perfectly Engaged", "Engaged", "idle<15s, Correct=95%, hints=0", "0.0-0.15", "Low", "None", "Green badge. Feedback: Excellent progress."],
    ["2. High Risk (Emotional)", "At-Risk", "idle>60s (unengaged), worsening SI", "0.6-0.8", "High", "Emotional", "Red badge. Feedback: Motivational check-in."],
    ["3. High Risk (Cognitive)", "At-Risk", "hints>1.5/q, SI>0.5", "0.5-0.7", "High", "Cognitive", "Red badge. Feedback: Content simplification."],
    ["4. High Risk (Behavioral)", "At-Risk", "attempts>4, idle>120s (total)", "0.5-0.8", "High", "Behavioral", "Red badge. Feedback: Suggest a break."],
    ["5. Med Risk (Recovering)", "Mixed", "SI improving, recovery events present", "0.2-0.4", "Medium", "None", "Yellow badge. Feedback: Resilience recognized."],
    ["6. Med Risk (Scaffolding)", "Mixed", "hints high but mastery improving", "0.2-0.4", "Medium", "Cognitive", "Yellow badge. Feedback: Fade hints slowly."],
    ["7. Low Risk (Occasional)", "Engaged", "occasional 3+ attempts, overall ok", "0.1-0.3", "Low", "None", "Green badge. Feedback: Normal learning curve."],
    ["8. Frustration Pattern", "At-Risk", "attempts>=3, idle<10s, incorrect", "0.4-0.6", "High", "Behavioral", "Red badge. Feedback: High frustration alert."],
    ["9. Gaming Pattern", "At-Risk", "High hints, low idle, low mastery", "0.3-0.5", "High", "Cognitive", "Red badge. Feedback: Hint abuse alert."],
    ["10. Cold Start", "New", "1st session, <5 rows", "N/A", "Low", "None", "Gray badge. Feedback: Insufficient data."],
    ["11. Module Switch", "Mixed", "Success in Logic, Failure in Maths", "0.2->0.7", "Medium", "Cognitive", "Yellow badge. Feedback: Domain-specific struggle."],
    ["12. Recovery Champion", "Mixed", "Many Struggling->Engaged transitions", "0.3-0.4", "Medium", "None", "Yellow badge. Feedback: Highly resilient."],
]
add_table(doc, headers, scenarios)
doc.add_page_break()

# SECTION 7
add_heading(doc, "Section 7 — Known Limitations and Future Work", 1)
doc.add_paragraph("• Discrepancy: Risk model assumes date-based sessions; Router uses 3-rule boundaries.")
doc.add_paragraph("• Limitation: Frustration is an approximation based on latency proxy.")
doc.add_paragraph("• BKT: Upgrade requires transition probability matrices (Learn/Slip/Guess).")
doc.add_paragraph("• Random Forest: Potential upgrade when cohort size > 50 students.")

add_footer(doc)
doc.save(OUT_PATH)
print(f"File saved to: {OUT_PATH}")
