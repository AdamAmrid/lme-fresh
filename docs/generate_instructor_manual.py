import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CONFIGURATION ---
OUT_FILE = "LME_Instructor_Analytics_Manual.docx"
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(DOCS_DIR, OUT_FILE)

# --- THEME COLORS ---
UM6P_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
UM6P_ORANGE = RGBColor(0xF4, 0x79, 0x20)
HEADER_BG = RGBColor(0x1E, 0x3A, 0x5F)
ROW_ALT_BG = RGBColor(0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LME Instructor Manual — Confidential · UM6P 2026")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = UM6P_BLUE
            run.font.size = Pt(16)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], HEADER_BG)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row().cells
        for c_idx, val in enumerate(row_data):
            row[c_idx].text = str(val)
            if r_idx % 2 == 1:
                set_cell_bg(row[c_idx], ROW_ALT_BG)
            p = row[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.font.size = Pt(9)
    return table

def add_callout(doc, text, label="INFO"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{label}: {text}")
    run.font.italic = True
    run.font.color.rgb = UM6P_ORANGE
    run.font.size = Pt(10)

# --- MAIN GENERATION ---
doc = Document()

# 1. TITLE PAGE
doc.add_paragraph("\n" * 5)
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("LME Instructor Analytics Manual")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = UM6P_BLUE

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = st.add_run("Deep Analysis & Intelligence Dashboard Guide")
run.font.size = Pt(16)
run.font.italic = True
run.font.color.rgb = UM6P_ORANGE

doc.add_paragraph("\n" * 10)
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("April 2026\nUniversité Mohammed VI Polytechnique")
run.font.size = Pt(12)

doc.add_page_break()

# 2. OVERVIEW
add_heading(doc, "Section 1: The Learner Modeling Engine (LME)", 1)
doc.add_paragraph(
    "The LME Dashboard provides real-time, AI-driven insights into student engagement. "
    "It uses behavioral telemetry to predict struggle and emotional frustration before they lead to attrition."
)

# 3. LEARNER PROFILES
add_heading(doc, "Section 2: Learner Profiles (Individual Analysis)", 1)
doc.add_paragraph("This view allows for a 'Deep Dive' into a single student's cognitive journey.")

headers = ["Component", "Description", "Analytical Value"]
rows = [
    ["Session Timeline", "Interactive path of circles (sessions).", "Identifies peaks of struggle and consistency."],
    ["Detailed Logic Table", "Every interaction (mistakes, hints, idle).", "Evidence-based discussion for parent meetings."],
    ["Risk Badge", "Numerical Score (0-1.00) and Label.", "Immediate quantification of intervention priority."],
    ["State Distribution", "Engaged vs Struggling vs Unengaged.", "Overall emotional profile of the student."]
]
add_table(doc, headers, rows)

# 4. COHORT OVERVIEW
add_heading(doc, "Section 3: Cohort Overview (Class-wide Analysis)", 1)
doc.add_paragraph("The bird's-eye view of your entire class cohort.")

add_heading(doc, "3.1 Engagement Distribution", 2)
doc.add_paragraph(
    "A bar chart showing the composition of states for every student. "
    "Features include a 'Class Average' baseline and an interactive slider bar for large classes."
)

add_heading(doc, "3.2 Session Heatmap", 2)
doc.add_paragraph("A high-density grid showing the 'Struggle Index' per session.")
add_callout(doc, "Vertical view shows student consistency. Horizontal view shows topic difficulty.")

add_heading(doc, "3.3 At-Risk Leaderboard", 2)
doc.add_paragraph("The prioritized queue of students requiring immediate intervention, ranked by their Risk Score.")

# 5. FORMULAS
add_heading(doc, "Section 4: The Intelligent Logic (Formulas)", 1)
doc.add_paragraph("The dashbard's intelligence is governed by these core formulas:")

add_heading(doc, "4.1 Struggle Index (SI)", 2)
doc.add_paragraph("SI = 0.4 * (Idle/120) + 0.4 * (Attempts/5) - 0.2 * (Score/1.0)")
add_callout(doc, "Lower is better. Values > 0.5 indicate significant stuckness.")

add_heading(doc, "4.2 Composite Risk Score", 2)
doc.add_paragraph("Risk = 30% Struggle + 30% Hint Dependency + 20% Unengagement + 20% Non-Participation")

add_heading(doc, "4.3 State Classification", 2)
rows = [
    ["Engaged", "Correct Answer OR (Attempts < 3 & Idle < 20s)"],
    ["Unengaged", "Idle time > 60 seconds"],
    ["Struggling", "Idle time >= 30s OR Attempts >= 3"],
    ["Engaged", "Otherwise (standard active state)"]
]
add_table(doc, ["State", "Trigger Logic"], rows)

add_footer(doc)
doc.save(OUT_PATH)
print(f"Manual generated successfully at: {OUT_PATH}")
